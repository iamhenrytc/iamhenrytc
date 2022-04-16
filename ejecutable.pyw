from faulthandler import disable
from os import stat
from tkinter import *
from tkinter import messagebox
from docx import Document

def main():
     
    #---------------------Funciones--------------------------------
    
    #Crear una nueva receta
    def NuevaReceta():
        ButBuscador.config(state=DISABLED)
        Butname.config(state=NORMAL)
        Lupa.config(state=DISABLED)
        nuevoname.config(state=NORMAL)
        nuevoname.delete(0, 'end')
        image =  messagebox.askquestion(message="¿Deseará agregar alguna imagen?", title="Importante")
        if image == "yes":
            Butfoto.config(state=NORMAL)
        else:
            Butfoto.config(state=DISABLED)


    #Guardar el titulo
    def GuardaTitulo():
        if myrecetanueva.get() == "":
            messagebox.showinfo(message="Ingrese un titulo aceptable.", title="Importante")
        else:
            messagebox.showinfo(message="Titulo asignado.", title="En hora buena")
            nuevoname.config(state=DISABLED)
            Butname.config(state=DISABLED)
            ButGuardar.config(state=NORMAL)
            Campo.config(state=NORMAL)
            Campo.delete("1.0","end")
            
    #Funcion para guardar los datos ingresados
    def GuardaDatos():
        nuevodocument = Document()
        nuevodocument.add_heading(myrecetanueva.get(), 0)
        #Agregamos la nueva receta a la lista 
        if Campo.get(1.0,END+"-1c") == "":
            messagebox.showinfo(message="Rellene el campo.", title="Importante")
        else:
            allRecetas = open("ListaRecetas.txt","a")
            allRecetas.write(f"{myrecetanueva.get()}\n")
            nuevodocument.add_paragraph(Campo.get(1.0,END+"-1c"))
            nuevodocument.save(f"{myrecetanueva.get()}.docx")
            messagebox.showinfo(message="Receta Guardada Exitosamente.", title="En hora buena")
            #Se procede a limpiar
            ButBuscador.config(state=NORMAL)
            Lupa.config(state=NORMAL)
            nuevoname.config(state=NORMAL)
            nuevoname.delete(0, 'end')
            nuevoname.config(state=DISABLED)
            Campo.delete("1.0","end")
            Campo.insert('1.0', 'Igrese los ingredientes :\n \nIgrese la preparacion de la receta :\n \nIgrese recomendaciones si las habria : ')
            Campo.config(state=DISABLED)
            ButGuardar.config(state=DISABLED)
            Butfoto.config(state=DISABLED)
    
    #Buscador de recetas
    #def Buscador():



    #---------------------Ventana y componentes-----------------------------------

    ventana = Tk()
    ventana.title('Las Recetas')
    ventana.configure(bg = "sky blue")
    ventana.resizable(0,0)
    #ventana.iconbitmap("name")
    
    myframe = Frame()
    myframe.pack(side="left",anchor="n")
    myframe.config(width="600",height="350",bg="sky blue")


    #Titulo del Programa 
    titulo = Label(myframe, text="Las Recetas: ",bg="sky blue",font=("Georgia", 20))
    titulo.place(x=10,y=15)

    #Buscador de Recetas
    myreceta=StringVar() 
    ButBuscador = Button(myframe, text="Buscar", command= lambda: Buscador())
    ButBuscador.place(x=530,y=20)
    Lupa = Entry(myframe,textvariable=myreceta)
    Lupa.place(x=400,y=20)

    #Creacion de nueva receta
    nueva = Button(myframe,text="Nueva Receta", command= lambda: NuevaReceta())
    nueva.place(x=180, y=20)

    #Titulo de receta nueva
    myrecetanueva=StringVar() 
    Butname = Button(myframe, text="Asignar Titulo", state=DISABLED,command= lambda: GuardaTitulo())
    Butname.place(x=330,y=50)
    nuevoname = Entry(myframe,textvariable=myrecetanueva, state=DISABLED)
    nuevoname.place(x=180,y=55)

    #Insertar la receta
    Campo=Text(myframe,width="50",height="15",highlightthickness=2)
    Campo.place(x=90, y=85)
    Campo.insert('1.0', 'Igrese los ingredientes :\n \nIgrese la preparacion de la receta :\n \nIgrese recomendaciones si las habria : ')
    Campo.config(state=DISABLED)
    #Guardar la receta 

    ButGuardar = Button(myframe, text="Guardar", state=DISABLED, command= lambda: GuardaDatos())
    ButGuardar.place(x=530,y=90)

    #Agregar foto de la receta
    Butfoto = Button(myframe, text="Insertar Foto", state=DISABLED)
    Butfoto.place(x=515,y=140)

    ventana.mainloop()

if __name__ == "__main__":
    main()